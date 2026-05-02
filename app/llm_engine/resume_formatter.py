"""
Resume formatting and export functionality.

Converts GeneratedResume objects to professional PDF and DOCX formats.
Handles text rendering, styling, and document generation.
"""

import io
import logging
import re
from typing import Any, Dict, List, Tuple, cast
from xml.sax.saxutils import escape

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
    KeepInFrame,
)
from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT
from reportlab.lib.colors import HexColor

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _normalize_url(value: str) -> str:
    """Normalize URL-like strings for safe hyperlink generation."""
    text = str(value or "").strip()
    if not text:
        return ""
    if text.startswith(("http://", "https://", "mailto:")):
        return text
    if "@" in text and " " not in text and "." in text:
        return f"mailto:{text}"
    return f"https://{text}"


def _link_markup(display: str, href: str) -> str:
    """Build safe reportlab link markup if URL is valid."""
    safe_display = escape(str(display or "").strip())
    safe_href = escape(_normalize_url(href))
    if not safe_display:
        return ""
    if not safe_href:
        return safe_display
    return f'<link href="{safe_href}">{safe_display}</link>'


def _add_docx_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Insert a clickable hyperlink run into a python-docx paragraph."""
    clean_text = str(text or "").strip()
    clean_url = _normalize_url(url)
    if not clean_text:
        return
    if not clean_url:
        paragraph.add_run(clean_text)
        return

    part = paragraph.part
    relationship_id = part.relate_to(clean_url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    new_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_props.append(run_style)

    new_run.append(run_props)
    text_elem = OxmlElement("w:t")
    text_elem.text = clean_text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _normalize_education_entries(education: Any) -> List[Tuple[str, str]]:
    """Convert education payload to structured title/details entries."""
    entries: List[Tuple[str, str]] = []

    if isinstance(education, list):
        for item in cast(List[Any], education):
            if isinstance(item, dict):
                institution = str(item.get("institution", "")).strip()
                degree = str(item.get("degree", "")).strip()
                duration = str(item.get("duration", "")).strip()
                title = institution or degree
                details = " | ".join([part for part in [degree if institution else "", duration] if part])
                if title or details:
                    entries.append((title or "Education", details))
            else:
                text = str(item).strip()
                if text:
                    entries.append((text, ""))
        return entries

    if isinstance(education, dict):
        institution = str(education.get("institution", "")).strip()
        degree = str(education.get("degree", "")).strip()
        duration = str(education.get("duration", "")).strip()
        title = institution or degree or "Education"
        details = " | ".join([part for part in [degree if institution else "", duration] if part])
        return [(title, details)]

    text = str(education or "").strip()
    if not text:
        return entries

    parts = [p.strip() for p in re.split(r"\n+|\s*;\s*", text) if p.strip()]
    for part in parts:
        split_pipe = [p.strip() for p in part.split("|", 1)]
        left = split_pipe[0]
        right = split_pipe[1] if len(split_pipe) > 1 else ""
        split_dash = [p.strip() for p in left.split(" - ", 1)]
        title = split_dash[0]
        detail_left = split_dash[1] if len(split_dash) > 1 else ""
        details = " | ".join([p for p in [detail_left, right] if p])
        entries.append((title or "Education", details))

    return entries


def _project_meta_line(project: Dict[str, Any]) -> str:
    """Build project metadata line with organization and optional date."""
    org = str(project.get("company", "")).strip() or str(project.get("name", "")).strip()
    duration = str(project.get("duration", "")).strip() or str(project.get("date", "")).strip() or str(project.get("dates", "")).strip()
    parts = [part for part in [org, duration] if part]
    return " | ".join(parts)


def _compact_for_one_page(resume: Dict[str, Any]) -> Dict[str, Any]:
    """Compress content density to maximize one-page fit while preserving entries."""
    compact = dict(resume)

    summary = str(compact.get("summary", "")).strip()
    if summary:
        words = summary.split()
        compact["summary"] = " ".join(words[:32]).strip()

    experience = compact.get("experience", [])
    if isinstance(experience, list):
        new_exp: List[Dict[str, Any]] = []
        for entry in experience:
            if not isinstance(entry, dict):
                continue
            copy = dict(entry)
            bullets = copy.get("bullets", [])
            if isinstance(bullets, list):
                clean_bullets = [str(b).strip() for b in bullets if str(b).strip()]
                copy["bullets"] = clean_bullets[:2] if clean_bullets else []
            new_exp.append(copy)
        compact["experience"] = new_exp

    projects = compact.get("projects", [])
    if isinstance(projects, list):
        new_projects: List[Dict[str, Any]] = []
        for entry in cast(List[Any], projects[:4]):
            if not isinstance(entry, dict):
                continue
            copy = dict(entry)
            if not str(copy.get("company", "")).strip():
                copy["company"] = str(copy.get("name", "")).strip()
            bullets = copy.get("bullets", [])
            if isinstance(bullets, list):
                clean_bullets = [str(b).strip() for b in bullets if str(b).strip()]
                copy["bullets"] = clean_bullets[:1] if clean_bullets else []
            new_projects.append(copy)
        compact["projects"] = _iter_unique_projects(new_projects)

    certs = compact.get("certifications", [])
    if isinstance(certs, list):
        compact["certifications"] = [str(c).strip() for c in certs if str(c).strip()][:3]

    return compact


def _experience_meta_line(exp: Dict[str, Any], fallback_location: str = "") -> str:
    """Build role/location/duration metadata line for experience entries."""
    title = str(exp.get("title", "")).strip()
    location = str(exp.get("location", "")).strip() or fallback_location
    duration = str(exp.get("duration", "")).strip()
    parts = [part for part in [title, location, duration] if part]
    return " | ".join(parts)


def _split_bullets_for_hierarchy(bullets: Any) -> Tuple[str, List[str]]:
    """Split bullets into one primary bullet and optional sub-bullets."""
    if not isinstance(bullets, list):
        return "", []
    cleaned = [str(item).strip() for item in cast(List[Any], bullets) if str(item).strip()]
    if not cleaned:
        return "", []
    if len(cleaned) >= 3:
        return cleaned[0], cleaned[1:]
    return "", cleaned


def _iter_unique_projects(projects: Any) -> List[Dict[str, Any]]:
    """Remove duplicate project titles while preserving first occurrence order."""
    if not isinstance(projects, list):
        return []

    unique_projects: List[Dict[str, Any]] = []
    seen_names: set[str] = set()
    for item in cast(List[Any], projects):
        if not isinstance(item, dict):
            continue
        project = cast(Dict[str, Any], item)
        raw_name = str(project.get("name", "")).strip()
        marker = re.sub(r"\s+", " ", raw_name).strip().lower() or "unnamed project"
        if marker in seen_names:
            continue
        seen_names.add(marker)
        unique_projects.append(project)

    return unique_projects


def _compact_skill_lines(skills_value: Any) -> List[str]:
    """Render compact recruiter-friendly skill lines with short labels."""
    label_map = {
        "programming_languages": "Languages",
        "data_science": "Data",
        "data_visualization": "Viz",
        "databases": "DB",
        "tools": "Tools",
    }
    lines: List[str] = []
    for category, skills in _iter_skill_categories(skills_value):
        short_label = label_map.get(category, category.replace("_", " ").title())
        lines.append(f"{short_label}: {', '.join(skills)}")
    return lines


def _iter_skill_categories(skills_value: Any) -> List[Tuple[str, List[str]]]:
    """Normalize arbitrary skills payload into typed category/value pairs."""
    if not isinstance(skills_value, dict):
        return []

    categories: List[Tuple[str, List[str]]] = []
    skills_map = cast(Dict[object, object], skills_value)
    for raw_category, raw_skills in skills_map.items():
        if not isinstance(raw_category, str) or not isinstance(raw_skills, list):
            continue

        raw_skill_list = cast(List[object], raw_skills)
        cleaned_skills = [str(skill).strip() for skill in raw_skill_list if str(skill).strip()]
        if cleaned_skills:
            categories.append((raw_category, cleaned_skills))
    return categories


def _build_docx_document(resume: Dict[str, Any], candidate_name: str) -> DocxDocument:
    """Build a DOCX document object for resume content."""
    resume = _compact_for_one_page(resume)
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)

    # HEADER - Candidate name
    if candidate_name:
        heading = doc.add_paragraph(candidate_name)
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # CONTACT INFO
    if resume.get("contact"):
        contact = cast(Dict[str, Any], resume["contact"])
        contact_line = doc.add_paragraph()
        contact_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def _append_sep() -> None:
            if contact_line.runs:
                sep = contact_line.add_run(" | ")
                sep.font.size = Pt(8)

        location = str(contact.get("location", "")).strip()
        if location:
            location_run = contact_line.add_run(location)
            location_run.font.size = Pt(8)

        email = str(contact.get("email", "")).strip()
        if email:
            _append_sep()
            _add_docx_hyperlink(contact_line, email, f"mailto:{email}")

        phone = str(contact.get("phone", "")).strip()
        if phone:
            _append_sep()
            phone_run = contact_line.add_run(phone)
            phone_run.font.size = Pt(8)

        linkedin = str(contact.get("linkedin", "")).strip()
        if linkedin:
            _append_sep()
            _add_docx_hyperlink(contact_line, linkedin, linkedin)

        github = str(contact.get("github", "")).strip()
        if github:
            _append_sep()
            _add_docx_hyperlink(contact_line, github, github)

    # SUMMARY
    if resume.get("summary"):
        doc.add_paragraph("PROFESSIONAL SUMMARY", style="Heading 2")
        doc.add_paragraph(resume["summary"])

    # EXPERIENCE
    if resume.get("experience"):
        doc.add_paragraph("PROFESSIONAL EXPERIENCE", style="Heading 2")
        fallback_location = ""
        if isinstance(resume.get("contact"), dict):
            fallback_location = str(cast(Dict[str, Any], resume.get("contact", {})).get("location", "")).strip()
        for exp in resume["experience"]:
            company = exp.get("company", "N/A")
            company_line = doc.add_paragraph(str(company))
            company_line.runs[0].font.bold = True
            company_line.runs[0].font.size = Pt(9)

            meta_line_text = _experience_meta_line(cast(Dict[str, Any], exp), fallback_location=fallback_location)
            if meta_line_text:
                meta_line = doc.add_paragraph(meta_line_text)
                meta_line.runs[0].font.size = Pt(8)
                meta_line.runs[0].italic = True

            main_bullet, sub_bullets = _split_bullets_for_hierarchy(exp.get("bullets", []))
            if main_bullet:
                doc.add_paragraph(main_bullet, style="List Bullet")
                for sub in sub_bullets:
                    sub_para = doc.add_paragraph(f"o {sub}")
                    sub_para.paragraph_format.left_indent = Inches(0.35)
                    for run in sub_para.runs:
                        run.font.size = Pt(9)
            else:
                for bullet in sub_bullets:
                    doc.add_paragraph(bullet, style="List Bullet")

    # PROJECTS
    if resume.get("projects"):
        doc.add_paragraph("PROJECT EXPERIENCE", style="Heading 2")
        for proj in _iter_unique_projects(resume.get("projects", [])):
            proj_name = proj.get("name", "Unnamed Project")

            p = doc.add_paragraph(proj_name)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9)

            meta = _project_meta_line(cast(Dict[str, Any], proj))
            if meta:
                meta_p = doc.add_paragraph(meta)
                meta_p.runs[0].font.italic = True
                meta_p.runs[0].font.size = Pt(8)

            for bullet in proj.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

            techs = proj.get("technologies", [])
            if techs:
                tech_para = doc.add_paragraph(f"Technologies: {', '.join(techs)}")
                for run in tech_para.runs:
                    run.font.italic = True
                    run.font.size = Pt(9)

    # EDUCATION
    if resume.get("education"):
        doc.add_paragraph("EDUCATION", style="Heading 2")
        for edu_title, edu_details in _normalize_education_entries(resume["education"]):
            edu_para = doc.add_paragraph()
            title_run = edu_para.add_run(edu_title)
            title_run.bold = True
            title_run.font.size = Pt(9)
            if edu_details:
                details_run = edu_para.add_run(f" | {edu_details}")
                details_run.font.size = Pt(8)

    # SKILLS
    if resume.get("skills"):
        doc.add_paragraph("TECHNICAL SKILLS", style="Heading 2")
        for line in _compact_skill_lines(resume.get("skills")):
            skills_para = doc.add_paragraph(line)
            skills_para.runs[0].font.bold = True

    # CERTIFICATIONS
    if resume.get("certifications"):
        doc.add_paragraph("CERTIFICATIONS", style="Heading 2")
        for cert in resume["certifications"]:
            doc.add_paragraph(cert, style="List Bullet")

    return doc


def render_resume_text(resume: Dict[str, Any]) -> str:
    """
    Convert GeneratedResume dict to formatted plain text.
    
    Returns clean text representation suitable for display or conversion.
    """
    resume = _compact_for_one_page(resume)
    lines: List[str] = []

    # CONTACT INFO
    if resume.get("contact"):
        contact = cast(Dict[str, Any], resume["contact"])
        contact_parts: List[str] = []
        if contact.get("location"):
            contact_parts.append(str(contact["location"]))
        if contact.get("email"):
            contact_parts.append(str(contact["email"]))
        if contact.get("phone"):
            contact_parts.append(str(contact["phone"]))
        if contact.get("linkedin"):
            contact_parts.append(str(contact["linkedin"]))
        if contact.get("github"):
            contact_parts.append(str(contact["github"]))
        if contact_parts:
            lines.append(" | ".join(contact_parts))
            lines.append("")

    # SUMMARY
    if resume.get("summary"):
        lines.append(resume["summary"])
        lines.append("")

    # EXPERIENCE
    if resume.get("experience"):
        lines.append("PROFESSIONAL EXPERIENCE")
        lines.append("=" * 80)
        fallback_location = ""
        if isinstance(resume.get("contact"), dict):
            fallback_location = str(cast(Dict[str, Any], resume.get("contact", {})).get("location", "")).strip()
        for exp in resume["experience"]:
            lines.append(str(exp.get("company", "N/A")))
            meta_line_text = _experience_meta_line(cast(Dict[str, Any], exp), fallback_location=fallback_location)
            if meta_line_text:
                lines.append(meta_line_text)

            main_bullet, sub_bullets = _split_bullets_for_hierarchy(exp.get("bullets", []))
            if main_bullet:
                lines.append(f"  • {main_bullet}")
                for sub in sub_bullets:
                    lines.append(f"    o {sub}")
            else:
                for bullet in sub_bullets:
                    lines.append(f"  • {bullet}")
            lines.append("")

    # PROJECTS
    if resume.get("projects"):
        lines.append("PROJECT EXPERIENCE")
        lines.append("=" * 80)
        for proj in _iter_unique_projects(resume.get("projects", [])):
            lines.append(str(proj.get("name", "Unnamed Project")))
            meta = _project_meta_line(cast(Dict[str, Any], proj))
            if meta:
                lines.append(meta)
            for bullet in proj.get("bullets", []):
                lines.append(f"  • {bullet}")
            
            techs = proj.get("technologies", [])
            if techs:
                lines.append(f"  Technologies: {', '.join(techs)}")
            lines.append("")

    # EDUCATION
    if resume.get("education"):
        lines.append("EDUCATION")
        lines.append("=" * 80)
        for edu_title, edu_details in _normalize_education_entries(resume["education"]):
            if edu_details:
                lines.append(f"{edu_title} | {edu_details}")
            else:
                lines.append(edu_title)
        lines.append("")

    # SKILLS
    if resume.get("skills"):
        lines.append("TECHNICAL SKILLS")
        lines.append("=" * 80)
        for line in _compact_skill_lines(resume.get("skills")):
            lines.append(f"  {line}")
        lines.append("")

    # CERTIFICATIONS
    if resume.get("certifications"):
        lines.append("CERTIFICATIONS")
        lines.append("=" * 80)
        for cert in resume["certifications"]:
            lines.append(f"  • {cert}")
        lines.append("")

    return "\n".join(lines)


def generate_pdf(
    resume: Dict[str, Any],
    output_path: str = "resume.pdf",
    candidate_name: str = "Candidate"
) -> str:
    """
    Generate a professional PDF resume from GeneratedResume dict.
    
    Args:
        resume: GeneratedResume dict with all sections
        output_path: Path where PDF will be saved
        candidate_name: Name to display in header
    
    Returns:
        Path to generated PDF
    """
    resume = _compact_for_one_page(resume)
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.35 * inch,
        leftMargin=0.35 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    header_style = ParagraphStyle(
        name="SectionHeader",
        parent=styles["Heading2"],
        fontSize=10,
        leading=11,
        spaceAfter=4,
        spaceBefore=5,
        textColor=HexColor("#000000"),
        fontName="Helvetica-Bold",
    )

    title_style = ParagraphStyle(
        name="JobTitle",
        parent=styles["Normal"],
        fontSize=9,
        leading=10,
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )

    company_style = ParagraphStyle(
        name="CompanyLine",
        parent=styles["Normal"],
        fontSize=8,
        leading=9,
        spaceAfter=2,
        fontName="Helvetica",
    )

    date_style = ParagraphStyle(
        name="DateLine",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceAfter=4,
        fontName="Helvetica",
        alignment=TA_RIGHT,
    )

    body_style = ParagraphStyle(
        name="Body",
        parent=styles["Normal"],
        fontSize=7.5,
        leading=8.5,
        spaceAfter=1,
        alignment=TA_JUSTIFY,
    )

    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=styles["Normal"],
        fontSize=7.5,
        leading=8.5,
        spaceAfter=1,
        leftIndent=0.14 * inch,
        bulletIndent=0.1 * inch,
    )

    sub_bullet_style = ParagraphStyle(
        name="SubBullet",
        parent=styles["Normal"],
        fontSize=7,
        leading=8,
        spaceAfter=1,
        leftIndent=0.26 * inch,
        bulletIndent=0.2 * inch,
    )

    content: List[Any] = []

    # HEADER - Candidate name
    if candidate_name:
        name_style = ParagraphStyle(
            name="CandidateName",
            fontSize=12,
            leading=13,
            spaceAfter=3,
            fontName="Helvetica-Bold",
            textColor=HexColor("#000000"),
        )
        content.append(Paragraph(candidate_name, name_style))

    # CONTACT INFO
    contact_style = ParagraphStyle(
        name="Contact",
        parent=styles["Normal"],
        fontSize=7,
        leading=8,
        spaceAfter=4,
        fontName="Helvetica",
    )
    
    if resume.get("contact"):
        contact = cast(Dict[str, Any], resume["contact"])
        contact_parts: List[str] = []
        if contact.get("location"):
            contact_parts.append(escape(str(contact["location"])))
        if contact.get("email"):
            contact_parts.append(_link_markup(str(contact["email"]), f"mailto:{str(contact['email'])}"))
        if contact.get("phone"):
            contact_parts.append(escape(str(contact["phone"])))
        if contact.get("linkedin"):
            contact_parts.append(_link_markup(str(contact["linkedin"]), str(contact["linkedin"])))
        if contact.get("github"):
            contact_parts.append(_link_markup(str(contact["github"]), str(contact["github"])))
        if contact_parts:
            contact_text = " | ".join([part for part in contact_parts if part])
            content.append(Paragraph(contact_text, contact_style))

    # SUMMARY
    if resume.get("summary"):
        content.append(Paragraph("PROFESSIONAL SUMMARY", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        content.append(Paragraph(resume["summary"], body_style))
        content.append(Spacer(1, 0.04 * inch))

    # EXPERIENCE
    if resume.get("experience"):
        content.append(Paragraph("PROFESSIONAL EXPERIENCE", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        fallback_location = ""
        if isinstance(resume.get("contact"), dict):
            fallback_location = str(cast(Dict[str, Any], resume.get("contact", {})).get("location", "")).strip()
        for exp in resume["experience"]:
            content.append(Paragraph(str(exp.get("company", "N/A")), title_style))
            meta_line_text = _experience_meta_line(cast(Dict[str, Any], exp), fallback_location=fallback_location)
            if meta_line_text:
                content.append(Paragraph(meta_line_text, company_style))

            main_bullet, sub_bullets = _split_bullets_for_hierarchy(exp.get("bullets", []))
            if main_bullet:
                content.append(Paragraph(f"• {main_bullet}", bullet_style))
                for sub in sub_bullets:
                    content.append(Paragraph(f"o {sub}", sub_bullet_style))
            else:
                for bullet in sub_bullets:
                    content.append(Paragraph(f"• {bullet}", bullet_style))

            content.append(Spacer(1, 0.03 * inch))

        content.append(Spacer(1, 0.04 * inch))

    # PROJECTS
    if resume.get("projects"):
        content.append(Paragraph("PROJECT EXPERIENCE", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for proj in _iter_unique_projects(resume.get("projects", [])):
            proj_name = proj.get("name", "Unnamed Project")
            content.append(Paragraph(f"<b>{proj_name}</b>", title_style))

            meta = _project_meta_line(cast(Dict[str, Any], proj))
            if meta:
                content.append(Paragraph(f"<i>{escape(meta)}</i>", company_style))

            for bullet in proj.get("bullets", []):
                content.append(Paragraph(f"• {bullet}", bullet_style))

            techs = proj.get("technologies", [])
            if techs:
                tech_text = f"<i>Technologies: {', '.join(techs)}</i>"
                content.append(Paragraph(tech_text, body_style))

            content.append(Spacer(1, 0.03 * inch))

        content.append(Spacer(1, 0.04 * inch))

    # EDUCATION
    if resume.get("education"):
        content.append(Paragraph("EDUCATION", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for edu_title, edu_details in _normalize_education_entries(resume["education"]):
            if edu_details:
                content.append(Paragraph(f"<b>{escape(edu_title)}</b> | {escape(edu_details)}", body_style))
            else:
                content.append(Paragraph(f"<b>{escape(edu_title)}</b>", body_style))
        content.append(Spacer(1, 0.04 * inch))

    # SKILLS
    if resume.get("skills"):
        content.append(Paragraph("TECHNICAL SKILLS", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for line in _compact_skill_lines(resume.get("skills")):
            skills_text = f"<b>{line.split(':', 1)[0]}:</b>{line.split(':', 1)[1] if ':' in line else ''}"
            content.append(Paragraph(skills_text, body_style))

        content.append(Spacer(1, 0.04 * inch))

    # CERTIFICATIONS
    if resume.get("certifications"):
        content.append(Paragraph("CERTIFICATIONS", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for cert in resume["certifications"]:
            content.append(Paragraph(f"• {cert}", bullet_style))

    try:
        frame_content: List[Any] = [KeepInFrame(maxWidth=7.8 * inch, maxHeight=10.5 * inch, content=content, mode="shrink")]
        doc.build(frame_content)
        logger.info(f"PDF generated successfully: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise


def generate_docx(
    resume: Dict[str, Any],
    output_path: str = "resume.docx",
    candidate_name: str = "Candidate"
) -> str:
    """
    Generate an editable DOCX resume from GeneratedResume dict.
    
    Allows users to further customize before submission.
    
    Args:
        resume: GeneratedResume dict with all sections
        output_path: Path where DOCX will be saved
        candidate_name: Name to display in header
    
    Returns:
        Path to generated DOCX
    """
    doc = _build_docx_document(resume, candidate_name)

    try:
        doc.save(output_path)
        logger.info(f"DOCX generated successfully: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        raise


def generate_docx_bytes(resume: Dict[str, Any], candidate_name: str = "Candidate") -> bytes:
    """Generate DOCX as bytes (for direct download without saving to disk)."""
    doc = _build_docx_document(resume, candidate_name)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def generate_pdf_bytes(resume: Dict[str, Any], candidate_name: str = "Candidate") -> bytes:
    """
    Generate PDF as bytes (for direct download without saving to disk).
    
    Useful for web apps that need to serve files directly.
    """
    resume = _compact_for_one_page(resume)
    output = io.BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=0.35 * inch,
        leftMargin=0.35 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
    )

    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        name="SectionHeader",
        parent=styles["Heading2"],
        fontSize=10,
        leading=11,
        spaceAfter=4,
        spaceBefore=5,
        textColor=HexColor("#000000"),
        fontName="Helvetica-Bold",
    )

    title_style = ParagraphStyle(
        name="JobTitle",
        parent=styles["Normal"],
        fontSize=9,
        leading=10,
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )

    company_style = ParagraphStyle(
        name="CompanyLine",
        parent=styles["Normal"],
        fontSize=8,
        leading=9,
        spaceAfter=2,
        fontName="Helvetica",
    )

    date_style = ParagraphStyle(
        name="DateLine",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceAfter=4,
        fontName="Helvetica",
        alignment=TA_RIGHT,
    )

    body_style = ParagraphStyle(
        name="Body",
        parent=styles["Normal"],
        fontSize=7.5,
        leading=8.5,
        spaceAfter=1,
        alignment=TA_JUSTIFY,
    )

    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=styles["Normal"],
        fontSize=7.5,
        leading=8.5,
        spaceAfter=1,
        leftIndent=0.14 * inch,
        bulletIndent=0.1 * inch,
    )

    sub_bullet_style = ParagraphStyle(
        name="SubBullet",
        parent=styles["Normal"],
        fontSize=7,
        leading=8,
        spaceAfter=1,
        leftIndent=0.26 * inch,
        bulletIndent=0.2 * inch,
    )

    content: List[Any] = []

    # HEADER
    if candidate_name:
        name_style = ParagraphStyle(
            name="CandidateName",
            fontSize=12,
            leading=13,
            spaceAfter=3,
            fontName="Helvetica-Bold",
            textColor=HexColor("#000000"),
        )
        content.append(Paragraph(candidate_name, name_style))

    # CONTACT INFO
    contact_style = ParagraphStyle(
        name="Contact",
        parent=styles["Normal"],
        fontSize=7,
        leading=8,
        spaceAfter=4,
        fontName="Helvetica",
    )
    
    if resume.get("contact"):
        contact = cast(Dict[str, Any], resume["contact"])
        contact_parts: List[str] = []
        if contact.get("location"):
            contact_parts.append(escape(str(contact["location"])))
        if contact.get("email"):
            contact_parts.append(_link_markup(str(contact["email"]), f"mailto:{str(contact['email'])}"))
        if contact.get("phone"):
            contact_parts.append(escape(str(contact["phone"])))
        if contact.get("linkedin"):
            contact_parts.append(_link_markup(str(contact["linkedin"]), str(contact["linkedin"])))
        if contact.get("github"):
            contact_parts.append(_link_markup(str(contact["github"]), str(contact["github"])))
        if contact_parts:
            contact_text = " | ".join([part for part in contact_parts if part])
            content.append(Paragraph(contact_text, contact_style))

    # SUMMARY
    if resume.get("summary"):
        content.append(Paragraph("PROFESSIONAL SUMMARY", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        content.append(Paragraph(resume["summary"], body_style))
        content.append(Spacer(1, 0.04 * inch))

    # EXPERIENCE
    if resume.get("experience"):
        content.append(Paragraph("PROFESSIONAL EXPERIENCE", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        fallback_location = ""
        if isinstance(resume.get("contact"), dict):
            fallback_location = str(cast(Dict[str, Any], resume.get("contact", {})).get("location", "")).strip()
        for exp in resume["experience"]:
            content.append(Paragraph(str(exp.get("company", "N/A")), title_style))
            meta_line_text = _experience_meta_line(cast(Dict[str, Any], exp), fallback_location=fallback_location)
            if meta_line_text:
                content.append(Paragraph(meta_line_text, company_style))

            main_bullet, sub_bullets = _split_bullets_for_hierarchy(exp.get("bullets", []))
            if main_bullet:
                content.append(Paragraph(f"• {main_bullet}", bullet_style))
                for sub in sub_bullets:
                    content.append(Paragraph(f"o {sub}", sub_bullet_style))
            else:
                for bullet in sub_bullets:
                    content.append(Paragraph(f"• {bullet}", bullet_style))

            content.append(Spacer(1, 0.03 * inch))

        content.append(Spacer(1, 0.04 * inch))

    # PROJECTS
    if resume.get("projects"):
        content.append(Paragraph("PROJECT EXPERIENCE", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for proj in _iter_unique_projects(resume.get("projects", [])):
            proj_name = proj.get("name", "Unnamed Project")
            content.append(Paragraph(f"<b>{proj_name}</b>", title_style))

            meta = _project_meta_line(cast(Dict[str, Any], proj))
            if meta:
                content.append(Paragraph(f"<i>{escape(meta)}</i>", company_style))

            for bullet in proj.get("bullets", []):
                content.append(Paragraph(f"• {bullet}", bullet_style))

            techs = proj.get("technologies", [])
            if techs:
                tech_text = f"<i>Technologies: {', '.join(techs)}</i>"
                content.append(Paragraph(tech_text, body_style))

            content.append(Spacer(1, 0.03 * inch))

        content.append(Spacer(1, 0.04 * inch))

    # EDUCATION
    if resume.get("education"):
        content.append(Paragraph("EDUCATION", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for edu_title, edu_details in _normalize_education_entries(resume["education"]):
            if edu_details:
                content.append(Paragraph(f"<b>{escape(edu_title)}</b> | {escape(edu_details)}", body_style))
            else:
                content.append(Paragraph(f"<b>{escape(edu_title)}</b>", body_style))
        content.append(Spacer(1, 0.04 * inch))

    # SKILLS
    if resume.get("skills"):
        content.append(Paragraph("TECHNICAL SKILLS", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for line in _compact_skill_lines(resume.get("skills")):
            skills_text = f"<b>{line.split(':', 1)[0]}:</b>{line.split(':', 1)[1] if ':' in line else ''}"
            content.append(Paragraph(skills_text, body_style))

        content.append(Spacer(1, 0.04 * inch))

    # CERTIFICATIONS
    if resume.get("certifications"):
        content.append(Paragraph("CERTIFICATIONS", header_style))
        content.append(HRFlowable(width="100%", thickness=0.6, color=HexColor("#666666"), spaceBefore=2, spaceAfter=6))
        for cert in resume["certifications"]:
            content.append(Paragraph(f"• {cert}", bullet_style))

    frame_content: List[Any] = [KeepInFrame(maxWidth=7.8 * inch, maxHeight=10.5 * inch, content=content, mode="shrink")]
    doc.build(frame_content)
    output.seek(0)
    return output.getvalue()


def generate_text_pdf_bytes(resume_text: str, candidate_name: str = "Candidate") -> bytes:
    """Generate professional PDF bytes from plain text resume content."""
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )

    styles = getSampleStyleSheet()

    # Name style
    name_style = ParagraphStyle(
        name="TextPDFName",
        parent=styles["Heading1"],
        fontSize=14,
        leading=16,
        spaceAfter=3,
        spaceBefore=0,
        textColor=HexColor("#000000"),
        fontName="Helvetica-Bold",
    )

    # Contact line style
    contact_style = ParagraphStyle(
        name="TextPDFContact",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceAfter=4,
        spaceBefore=0,
        fontName="Helvetica",
    )

    # Section header style
    section_style = ParagraphStyle(
        name="TextPDFSection",
        parent=styles["Normal"],
        fontSize=11,
        leading=12,
        spaceBefore=6,
        spaceAfter=2,
        fontName="Helvetica-Bold",
        textColor=HexColor("#000000"),
    )

    # Company or role line style
    company_style = ParagraphStyle(
        name="TextPDFCompany",
        parent=styles["Normal"],
        fontSize=10,
        leading=12,
        spaceBefore=2,
        spaceAfter=1,
        fontName="Helvetica-Bold",
    )

    # Regular body line
    job_style = ParagraphStyle(
        name="TextPDFJob",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=1,
        fontName="Helvetica",
    )

    # Bullet style
    bullet_style = ParagraphStyle(
        name="TextPDFBullet",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=2,
        fontName="Helvetica",
        leftIndent=0.15 * inch,
        bulletIndent=0.1 * inch,
    )

    # Sub-bullet style
    subbullet_style = ParagraphStyle(
        name="TextPDFSubBullet",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceBefore=0,
        spaceAfter=2,
        fontName="Helvetica",
        leftIndent=0.25 * inch,
        bulletIndent=0.1 * inch,
    )

    def _is_contact_line(value: str) -> bool:
        return "|" in value or ("@" in value and "." in value)

    def _is_section_header(value: str) -> bool:
        upper = value.isupper()
        return upper and 2 <= len(value) <= 45 and any(
            kw in value
            for kw in [
                "SUMMARY",
                "EXPERIENCE",
                "PROJECT",
                "EDUCATION",
                "SKILL",
                "CERTIFICATION",
                "ACHIEVEMENT",
                "AWARD",
                "LANGUAGE",
            ]
        )

    def _is_bullet_line(value: str) -> bool:
        return value.startswith("•") or value.startswith("-") or value.startswith("–")

    content: List[Any] = []
    lines = [ln.strip() for ln in str(resume_text or "").splitlines()]
    lines = [ln for ln in lines if ln]

    if not lines:
        if candidate_name:
            content.append(Paragraph(escape(candidate_name), name_style))
        doc.build(content)
        output.seek(0)
        return output.getvalue()

    first_line = lines[0]
    candidate_display_name = candidate_name.strip() if candidate_name else ""
    if not candidate_display_name or candidate_display_name.lower() == "candidate":
        candidate_display_name = first_line

    content.append(Paragraph(escape(candidate_display_name), name_style))

    start_idx = 0
    if first_line.lower() == candidate_display_name.lower():
        start_idx = 1

    for idx in range(start_idx, len(lines)):
        line = lines[idx]

        if _is_contact_line(line):
            content.append(Paragraph(escape(line), contact_style))
            continue

        if _is_section_header(line):
            content.append(Paragraph(escape(line), section_style))
            # Keep explicit section separator lines in PDF.
            content.append(
                HRFlowable(
                    width="100%",
                    thickness=0.7,
                    color=HexColor("#444444"),
                    spaceBefore=1,
                    spaceAfter=3,
                )
            )
            continue

        if _is_bullet_line(line):
            bullet_text = line.lstrip("•-–").strip()
            content.append(Paragraph(f"• {escape(bullet_text)}", bullet_style))
            continue

        if line.startswith("o") or line.startswith("○"):
            sub_bullet_text = line.lstrip("o○").strip()
            content.append(Paragraph(f"○ {escape(sub_bullet_text)}", subbullet_style))
            continue

        has_date_tokens = any(ch.isdigit() for ch in line) and ("-" in line or "–" in line)
        if has_date_tokens:
            content.append(Paragraph(escape(line), company_style))
            continue

        content.append(Paragraph(escape(line), job_style))

    doc.build(content)
    output.seek(0)
    return output.getvalue()


def generate_text_docx_bytes(resume_text: str, candidate_name: str = "Candidate") -> bytes:
    """Generate professional DOCX bytes from plain text resume content."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def _is_contact_line(value: str) -> bool:
        return "|" in value or ("@" in value and "." in value)

    def _is_section_header(value: str) -> bool:
        upper = value.isupper()
        return upper and 2 <= len(value) <= 45 and any(
            kw in value
            for kw in [
                "SUMMARY",
                "EXPERIENCE",
                "PROJECT",
                "EDUCATION",
                "SKILL",
                "CERTIFICATION",
                "ACHIEVEMENT",
                "AWARD",
                "LANGUAGE",
            ]
        )

    def _is_bullet_line(value: str) -> bool:
        return value.startswith("•") or value.startswith("-") or value.startswith("–")

    lines = [ln.strip() for ln in str(resume_text or "").splitlines()]
    lines = [ln for ln in lines if ln]

    if not lines:
        if candidate_name:
            heading = doc.add_paragraph(candidate_name)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if heading.runs:
                heading.runs[0].font.size = Pt(14)
                heading.runs[0].font.bold = True
                heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    first_line = lines[0]
    candidate_display_name = candidate_name.strip() if candidate_name else ""
    if not candidate_display_name or candidate_display_name.lower() == "candidate":
        candidate_display_name = first_line

    heading = doc.add_paragraph(candidate_display_name)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if heading.runs:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.paragraph_format.space_after = Pt(2)
    heading.paragraph_format.space_before = Pt(0)

    start_idx = 0
    if first_line.lower() == candidate_display_name.lower():
        start_idx = 1

    for idx in range(start_idx, len(lines)):
        line = lines[idx]

        if _is_contact_line(line):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            continue

        if _is_section_header(line):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)

            # Keep explicit section separator lines in DOCX.
            separator = doc.add_paragraph("_" * 90)
            for run in separator.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(68, 68, 68)
            separator.paragraph_format.space_before = Pt(0)
            separator.paragraph_format.space_after = Pt(3)
            continue

        if _is_bullet_line(line):
            bullet_text = line.lstrip("•-–").strip()
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            continue

        if line.startswith("o") or line.startswith("○"):
            sub_bullet_text = line.lstrip("o○").strip()
            p = doc.add_paragraph(sub_bullet_text, style="List Bullet 2")
            for run in p.runs:
                run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            continue

        has_date_tokens = any(ch.isdigit() for ch in line) and ("-" in line or "–" in line)
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(10 if has_date_tokens else 9)
            run.font.bold = bool(has_date_tokens)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
